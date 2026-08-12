from __future__ import annotations

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output"
FIG = OUT / "figures"
OUT.mkdir(exist_ok=True)
FIG.mkdir(exist_ok=True)

# narrative_proposal preset + NAR manuscript override
TOKENS = {
    "page_width": 8.5,
    "page_height": 11.0,
    "margin": 1.0,
    "body_font": "Times New Roman",
    "body_size": 12,
    "body_color": "111111",
    "heading_color": "17365D",
    "accent": "1F6E8C",
    "line_spacing": 1.0,
    "space_after": 6,
}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text: str, url: str, color="0563C1"):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def _font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            pass
    return ImageFont.load_default()


def _centered(draw, box, text, font, fill, spacing=5):
    x0,y0,x1,y1=box
    bounds=draw.multiline_textbbox((0,0),text,font=font,align="center",spacing=spacing)
    w,h=bounds[2]-bounds[0],bounds[3]-bounds[1]
    draw.multiline_text(((x0+x1-w)/2,(y0+y1-h)/2),text,font=font,fill=fill,align="center",spacing=spacing)


def _arrow(draw, start, end, color="#1F6E8C", width=5):
    draw.line([start,end],fill=color,width=width)
    ex,ey=end; sx,sy=start
    import math
    a=math.atan2(ey-sy,ex-sx)
    l=18
    pts=[(ex,ey),(ex-l*math.cos(a-.55),ey-l*math.sin(a-.55)),(ex-l*math.cos(a+.55),ey-l*math.sin(a+.55))]
    draw.polygon(pts,fill=color)


def make_figures():
    # Graphical abstract / architecture
    im=Image.new("RGB",(2100,1040),"white"); draw=ImageDraw.Draw(im)
    boxes = [
        (50, 160, 460, 390, "Public resources", "UniProt · KinHub · GTEx\nClinVar · PDB · ChEMBL\nPubChem · STRING · PubMed"),
        (610, 160, 1020, 390, "Evidence-aware ETL", "Reconcile identifiers\nValidate provenance\nRecord source metadata"),
        (1170, 160, 1580, 390, "MongoDB", "678 catalogue entries\n157,342 audited records\nTyped collections"),
        (1730, 160, 2050, 390, "Web APIs", "Validated filters\nPagination\nSource links"),
        (180, 620, 650, 860, "Explore", "Search · group · organ\nPDIS histogram · tree"),
        (815, 620, 1285, 860, "Inspect", "Kinase dossiers\nStructure · expression\nvariants · associations"),
        (1450, 620, 1920, 860, "Ask", "Evidence-grounded assistant\nKinomeX + source databases\nPMID/DOI for literature"),
    ]
    for x0,y0,x1,y1,title,body in boxes:
        fill="#EAF4F8" if y0<500 else "#F3F0FA"
        draw.rounded_rectangle((x0,y0,x1,y1),radius=24,fill=fill,outline="#1F6E8C",width=4)
        _centered(draw,(x0,y0+25,x1,y0+105),title,_font(31,True),"#17365D")
        _centered(draw,(x0,y0+95,x1,y1-18),body,_font(22),"#243442",spacing=8)
    for s,e in [((460,275),(610,275)),((1020,275),(1170,275)),((1580,275),(1730,275))]: _arrow(draw,s,e)
    for e in [(415,620),(1050,620),(1685,620)]: _arrow(draw,(1375,390),e,color="#7A5AA6",width=4)
    _centered(draw,(100,20,2000,125),"KinomeX: source-linked integration and exploration of the human kinase landscape",_font(38,True),"#17365D")
    im.save(FIG/"graphical_abstract.png",dpi=(300,300))

    groups = {"AGC":64,"CAMK":76,"CK1":12,"CMGC":65,"STE":47,"TK":91,"TKL":43,"RGC":5,"Atypical":194,"Other":81}
    im=Image.new("RGB",(1900,1100),"white"); draw=ImageDraw.Draw(im)
    names=list(groups); vals=list(groups.values())
    colors = ["#35B8E0","#8B5FBF","#E7A42B","#3ABF9B","#DF526C","#438CE1","#E6792E","#26A6A1","#8794A6","#737373"]
    left,top,right,bottom=150,180,1840,860
    draw.text((150,45),"Reconciled KinomeX catalogue by displayed group",font=_font(38,True),fill="#17365D")
    for tick in [0,50,100,150,200]:
        y=bottom-(tick/200)*(bottom-top)
        draw.line((left,y,right,y),fill="#D8DEE3",width=2)
        draw.text((75,y-15),str(tick),font=_font(22),fill="#384B59")
    bw=115; gap=50
    for i,(name,v,color) in enumerate(zip(names,vals,colors)):
        x=left+i*(bw+gap)+20; h=(v/200)*(bottom-top); y=bottom-h
        draw.rounded_rectangle((x,y,x+bw,bottom),radius=10,fill=color)
        _centered(draw,(x,y-45,x+bw,y-5),str(v),_font(23,True),"#243442")
        _centered(draw,(x,bottom+18,x+bw,bottom+65),name,_font(21),"#243442")
    draw.text((150,945),"678 protein entries = 522 KinHub-indexed core entries + 156 reviewed UniProt extensions.",font=_font(24),fill="#384B59")
    draw.text((150,990),"The historical 536 count refers to kinase-domain rows, not distinct protein entries.",font=_font(24),fill="#384B59")
    im.save(FIG/"catalogue_groups.png",dpi=(300,300))


def configure_document(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Inches(TOKENS["page_width"])
    sec.page_height = Inches(TOKENS["page_height"])
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(TOKENS["margin"])
    sec.header_distance = Inches(0.45); sec.footer_distance = Inches(0.45)
    add_page_number(sec.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = TOKENS["body_font"]
    normal.font.size = Pt(TOKENS["body_size"])
    normal.font.color.rgb = rgb(TOKENS["body_color"])
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), TOKENS["body_font"])
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = TOKENS["line_spacing"]
    normal.paragraph_format.space_after = Pt(TOKENS["space_after"])
    normal.paragraph_format.widow_control = True

    for style_name, size, before, after in [("Title",18,0,12),("Heading 1",14,14,6),("Heading 2",12,10,4),("Heading 3",11,8,3)]:
        style = doc.styles[style_name]
        style.font.name = TOKENS["body_font"]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(TOKENS["heading_color"] if style_name != "Title" else "111111")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), TOKENS["body_font"])
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Figure Caption" not in doc.styles:
        s = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else: s = doc.styles["Figure Caption"]
    s.font.name = TOKENS["body_font"]; s.font.size = Pt(10)
    s.paragraph_format.space_before = Pt(4); s.paragraph_format.space_after = Pt(8)
    s.paragraph_format.keep_with_next = True

    for name in ["Reference", "Compact"]:
        if name not in doc.styles: doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    doc.styles["Reference"].font.name = TOKENS["body_font"]
    doc.styles["Reference"].font.size = Pt(10)
    doc.styles["Reference"].paragraph_format.left_indent = Inches(.25)
    doc.styles["Reference"].paragraph_format.first_line_indent = Inches(-.25)
    doc.styles["Reference"].paragraph_format.space_after = Pt(3)
    doc.styles["Compact"].font.name = TOKENS["body_font"]
    doc.styles["Compact"].font.size = Pt(10)
    doc.styles["Compact"].paragraph_format.space_after = Pt(2)


def add_para(doc, text="", style=None, align=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else: p.add_run(text)
    if align is not None: p.alignment = align
    return p


def add_source_table(doc):
    headers = ["Source", "Information used", "Current records/coverage", "Reference"]
    rows = [
        ("UniProtKB/Swiss-Prot", "Reviewed identity, sequence, domains, function, catalytic activity and disease comments", "678 catalogue entries; 256 entries with disease comments", "(3)"),
        ("KinHub/Manning", "Core roster, domain rows, group/family classification", "522 protein entries; 536 kinase-domain rows", "(1,2)"),
        ("GTEx v10", "Median tissue expression and organ mapping", "36,396 gene–tissue records; 674 genes", "(7)"),
        ("ClinVar", "Clinically interpreted variants", "63,649 records; 260 genes", "(8)"),
        ("RCSB PDB", "Experimentally determined structures and resolution", "10,034 kinase-linked structures; 476 genes", "(5)"),
        ("ChEMBL/PubChem", "Target-linked compounds and assay records", "45,646 records; 43,561 ChEMBL and 2,085 PubChem", "(9,10)"),
        ("UniProt disease comments", "Curated disease associations and identifiers", "256 gene-level documents", "(3)"),
        ("STRING", "Live human functional/physical association retrieval", "Query dependent; confidence and evidence channels retained", "(4)"),
        ("PubMed/ClinicalTrials.gov", "Literature and study counts used in PDIS; live referenced evidence", "Query dependent; 678 PDIS records", "(11,12)"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [1.25, 2.55, 2.15, .65]
    for i,(h,w) in enumerate(zip(headers,widths)):
        c=table.rows[0].cells[i]; c.width=Inches(w); set_cell_shading(c,"D9EAF2"); set_cell_margins(c)
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9)
    for row in rows:
        cells=table.add_row().cells
        for i,(value,w) in enumerate(zip(row,widths)):
            cells[i].width=Inches(w); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP; set_cell_margins(cells[i])
            cells[i].paragraphs[0].style="Compact"; cells[i].paragraphs[0].add_run(value)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    return table


def build():
    make_figures()
    doc = Document()
    configure_document(doc)
    cp = doc.core_properties
    cp.title = "KinomeX: an evidence-traceable research server for the human protein kinase landscape"
    cp.subject = "Draft manuscript for Nucleic Acids Research Web Server Issue"
    cp.author = "[Authors to be finalized]"
    cp.keywords = "human kinome; protein kinase; database; web server; PDIS; evidence provenance"

    p=doc.add_paragraph(style="Title"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(cp.title)
    add_para(doc,"[Author names to be finalized], Nikolay V. Dokholyan*",align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,"Department of Pharmacology, University of Virginia School of Medicine, Charlottesville, VA 22908, USA",align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,"*To whom correspondence should be addressed. Email: [corresponding-author email]",align=WD_ALIGN_PARAGRAPH.CENTER)
    p=add_para(doc,"Manuscript type: Web Server article | Draft date: 11 August 2026",style="Compact",align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].italic=True

    doc.add_heading("ABSTRACT", level=1)
    abstract=("Protein kinase information is distributed among resources that use different identifiers, coverage criteria and evidence models. KinomeX is a web-based research environment that reconciles a KinHub-indexed human kinase core with reviewed UniProtKB/Swiss-Prot protein kinase entries and connects each catalogue record to source-linked molecular and biomedical evidence. The audited database contains 678 protein entries, comprising 522 KinHub-indexed entries and 156 reviewed UniProt extensions; these represent 536 KinHub kinase-domain rows. Kinase dossiers integrate reviewed functional annotation, 10,034 experimentally determined structures, 45,646 compound or assay records, 36,396 tissue-expression records, 63,649 ClinVar variants and 256 gene-level disease records. The Explorer supports inclusive group, organ and Pharmaceutical Development Interest Score (PDIS) filters, pagination-aware totals and an interactive score distribution. A radial kinome tree, molecular viewers and live STRING association retrieval provide complementary structural and network views. An optional research assistant retrieves KinomeX records and connected source evidence, links database-derived statements to their originating records and requires matched PMID and DOI identifiers for literature-derived claims. KinomeX is intended for evidence inspection and hypothesis generation; PDIS summarizes development-associated evidence and is not a measure of biological importance or clinical utility. The server is available at [PUBLIC HTTPS URL TO BE INSERTED].")
    add_para(doc,abstract)
    add_para(doc,"Keywords: human kinome; protein kinase; evidence integration; web server; molecular annotation; target exploration",style="Compact")

    doc.add_heading("GRAPHICAL ABSTRACT", level=1)
    doc.add_picture(str(FIG/"graphical_abstract.png"), width=Inches(6.45))
    add_para(doc,"Graphical abstract. KinomeX reconciles kinase identity and classification, preserves source provenance during ingestion, and presents catalogue exploration, kinase-level evidence and source-aware conversational retrieval through a single web environment.",style="Figure Caption")

    doc.add_heading("INTRODUCTION", level=1)
    for text in [
        "Protein kinases regulate cellular processes by transferring phosphate groups to protein and non-protein substrates. The human protein kinase complement was organized into a widely used group and family framework by Manning and colleagues (1). Subsequent resources have provided browsable mappings of this classification (2), while reviewed protein records, structures, molecular associations, bioactivity measurements, tissue expression and clinically interpreted variants are maintained by separate databases (3–10). These resources differ in unit of representation: a kinase-domain catalogue can contain multiple rows for one protein, whereas a protein knowledgebase generally uses one accession per protein entry. Counts therefore require an explicit definition.",
        "A researcher examining one kinase commonly moves among multiple websites to establish identity, read reviewed functional annotation, locate structures, inspect tissue expression, review reported variants and evaluate molecular associations. Directly merging these records without provenance can obscure differences in source scope and evidence type. A functional association in STRING, for example, is not equivalent to evidence of direct physical binding (4). Likewise, a ClinVar clinical-significance assertion reflects the submitted and reviewed record rather than an independent KinomeX interpretation (8).",
        "We developed KinomeX to provide a reconciled, source-linked view of the human kinase landscape. The resource combines a defined catalogue, kinase dossiers, an Explorer with composable filters, a radial kinase tree, structural and network visualizations and an optional evidence-grounded research assistant. This report describes the catalogue definition, ingestion and validation procedures, current database contents, PDIS calculation, interface and evidence-control rules. The prose and interface distinguish source observations from derived quantities and explicitly display absent records as unavailable."
    ]: add_para(doc,text)

    doc.add_heading("MATERIALS AND METHODS", level=1)
    doc.add_heading("Catalogue definition and identifier reconciliation", level=2)
    add_para(doc,"KinomeX defines a catalogue entry as a protein-level record anchored to a UniProt accession and gene symbol. The core partition contains UniProt entries represented by at least one row in the KinHub roster; the extension partition contains reviewed human UniProt entries carrying the controlled Protein kinase keyword (KW-0418) that are not represented in the core. KinHub classification traces to the Manning human kinome framework (1,2), and protein metadata are retrieved from reviewed UniProtKB/Swiss-Prot records (3). The current reconciliation comprises 522 core entries and 156 extensions. The 522 core entries represent 536 kinase-domain rows because some proteins contain more than one annotated kinase domain. Of 625 reviewed human UniProt entries carrying KW-0418, 469 overlap the KinHub roster; 53 additional KinHub accessions resolve through UniProt outside this keyword set. PRKY/O43930 is retained as a labeled historical/inactive record. Catalogue metadata stores these definitions, source URLs, retrieval time and partition counts.")

    doc.add_heading("Data acquisition and provenance", level=2)
    add_para(doc,"An asynchronous Python pipeline performs idempotent, kinase-scoped retrieval and writes normalized documents to MongoDB. Each imported record retains a source label and the identifiers needed to return to the originating resource. Reviewed function, catalytic activity, sequence, domains and disease comments are obtained from UniProtKB/Swiss-Prot (3). Human protein associations are queried from STRING with species identifier 9606 and retain combined, experimental, database and text-mining scores (4). Experimentally determined structures are retrieved from the RCSB Protein Data Bank (5). AlphaFold Database links can supplement profile navigation but predicted models are not counted as experimentally determined structures (6). Median tissue expression is obtained from GTEx v10 (7); clinically interpreted variants are obtained from ClinVar (8); and kinase-linked compound and assay records are obtained from ChEMBL and PubChem (9,10). Literature retrieval uses PubMed/NCBI E-utilities (11), and study counts are queried from ClinicalTrials.gov (12).")
    add_source_table(doc)
    add_para(doc,"Table 1. Data sources and the audited database snapshot used for this manuscript (11 August 2026). Counts describe records stored in the local KinomeX instance and are not counts of unique biological observations across source databases.",style="Figure Caption")

    doc.add_heading("Database audit and quality controls", level=2)
    add_para(doc,"The manuscript snapshot was generated by non-mutating catalogue and database audits. The database contained 157,342 documents, no records matching the project’s synthetic-development signatures and no non-finite numeric values. Required catalogue identifiers and membership labels were present for all 678 entries. Expression, variant, disease, PDIS and structure collections contained no orphan gene symbols. Three historical gene symbols in the bioactivity collection (CDC2L2, PCTK1 and TSSK1) did not match current catalogue symbols and are reported as unresolved rather than silently reassigned. Source-specific record counts and retrieval dates are preserved in audit metadata.")

    doc.add_heading("Pharmaceutical Development Interest Score", level=2)
    add_para(doc,"PDIS is a derived evidence-density measure intended to summarize the amount of pharmaceutical-development information associated with a kinase. Four components are calculated on a 0–100 scale: PubMed citation count, selected active or completed ClinicalTrials.gov study count, experimentally determined structure coverage and distinct compound count. Citation and compound components are logarithmically normalized to the largest count in the current catalogue. The clinical-trial component is capped at 100 after normalization to 100 studies. The structure component combines the best-resolution score (60%) and average-resolution score (40%) over an approximately 1.5–4.0 Å interval. The implemented total is:")
    p=add_para(doc,"PDIS = [0.30(citation) + 0.30(clinical trials) + 0.15(structure) + 0.15(compound diversity)] / 0.90",align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].bold=True
    add_para(doc,"All four prerequisites must pass source and kinase-scope validation before calculation. Missing PDIS documents are represented as unavailable and are not converted to zero. Stored totals range from 0 to 95.74 in the current snapshot. PDIS does not estimate kinase function, biological importance, treatment efficacy, safety or clinical priority.")

    doc.add_heading("Web application and API", level=2)
    add_para(doc,"KinomeX is implemented with the Next.js App Router, React and TypeScript. Server-side API routes validate input, construct MongoDB queries and enrich results from normalized collections. List, histogram and selected aggregate responses use bounded process-local caching. The principal kinase endpoint supports escaped gene/name search, displayed group, organ system, inclusive PDIS interval, catalogue partition, sort order and bounded pagination. Organ and PDIS conditions are resolved to gene sets and intersected before total counting; consequently, page totals represent the complete filtered result rather than the current page. Kinases without a PDIS document are returned when no score interval is requested and excluded when an explicit interval is supplied.")

    doc.add_heading("Evidence-grounded assistant", level=2)
    add_para(doc,"The optional assistant combines deterministic query parsing with retrieval from KinomeX, reviewed UniProt annotations, STRING associations and PubMed. Database claims are linked to the supplying record. Literature-derived statements are emitted only when the retrieved article has both a PubMed identifier and a matching DOI; each statement or table row includes those identifiers and links. STRING-derived rows link the specific human association or network view and describe the score as association confidence. The assistant does not substitute model memory for unavailable source evidence. Conversation state is stored for the browser session, allowing navigation among pages without discarding the research conversation. The model endpoint is configurable through an OpenAI-compatible chat-completions interface; the resource can be used without enabling the assistant.")

    doc.add_heading("RESULTS", level=1)
    doc.add_heading("Current catalogue and evidence coverage", level=2)
    add_para(doc,"The reconciled catalogue contains 678 distinct accessions and 678 distinct gene symbols. Displayed group counts are AGC, 64; CAMK, 76; CK1, 12; CMGC, 65; STE, 47; TK, 91; TKL, 43; RGC, 5; Atypical, 194; and Other, 81. The Atypical and Other categories are larger than their KinHub core partitions because reviewed UniProt extensions are retained and assigned to displayed categories rather than omitted. Figure 1 reports protein entries, not kinase domains.")
    doc.add_picture(str(FIG/"catalogue_groups.png"), width=Inches(6.45))
    add_para(doc,"Figure 1. Reconciled KinomeX catalogue by displayed group. Counts sum to 678 protein entries. The catalogue contains 522 KinHub-indexed core entries and 156 reviewed UniProt extensions; the core represents 536 kinase-domain rows.",style="Figure Caption")
    add_para(doc,"The audited evidence layer contains expression records for 674 catalogue genes, ClinVar records for 260 genes, UniProt disease documents for 256 genes and experimentally determined structures for 476 genes. All 678 catalogue entries have a current PDIS document because the required imports passed the implemented validation gates. Record counts describe the stored snapshot and can change when upstream databases or the local ETL snapshot change.")

    doc.add_heading("Catalogue exploration", level=2)
    add_para(doc,"The Explorer applies group, organ, name and score filters to a common result set. A 20-bin histogram shows the PDIS distribution for the active name, group and organ context while retaining the complete score domain for adjustment. Slider boundaries are inclusive and preserve fractional values without display rounding. The sidebar reports the filtered total and group breakdown from the complete query result. Pagination changes the displayed subset but not the total. These behaviors were tested for PDIS boundary values, absent PDIS records, organ-plus-PDIS intersections and page totals.")
    add_para(doc,"The radial tree provides an orthogonal group-oriented view. Labels are progressively disclosed as the user zooms; collision handling and scale-aware label placement maintain legibility at deep zoom. Selecting a node opens its kinase dossier.")

    doc.add_heading("Kinase dossiers", level=2)
    add_para(doc,"Each dossier begins with a summary, followed by the reviewed UniProtKB/Swiss-Prot function text when available. Separate views present structures, tissue distribution, ligand bioactivity, ClinVar variants, STRING associations, disease annotations and source references. Structure panels link PDB entries and render supported coordinates in interactive molecular viewers. Expression panels retain tissue names and median TPM values from GTEx rather than inferring expression from organ labels. Variant and disease panels preserve ClinVar and UniProt source semantics. The network view queries STRING at use time, allows functional or physical association mode and confidence adjustment, and reports component evidence scores. Association edges are not labeled as direct binding unless the source evidence supports that interpretation.")

    doc.add_heading("Source-aware research conversation", level=2)
    add_para(doc,"Natural-language requests may combine catalogue filters and external evidence. When a query can be answered from structured KinomeX data, the assistant returns linked kinase rows with stored fields. Functional questions may retrieve reviewed UniProt text and STRING associations. Literature questions query PubMed and return only articles for which both PMID and DOI are verified. This design allows, for example, a kinase–cytoskeleton association supported by STRING to be cited to STRING without requiring a literature identifier, while a mechanistic statement derived from an article must carry its PMID and DOI. The interface displays source labels and links with the answer rather than presenting an undifferentiated response.")

    doc.add_heading("DISCUSSION", level=1)
    add_para(doc,"KinomeX addresses two recurrent sources of ambiguity in kinase-resource use: catalogue accounting and evidence provenance. Protein-level and domain-level counts are reported separately, and the extension beyond the historical core is defined by a reproducible reviewed-UniProt query. The resulting catalogue is broader than a conventional protein-kinase-domain census and should be interpreted as an accounted set of protein entries that combines the KinHub core with reviewed UniProt extensions.")
    add_para(doc,"The resource is designed for inspection and hypothesis generation. It does not replace the source databases, primary literature or expert review. Upstream records can change independently, live services can be unavailable and database assertions can reflect different evidence standards. PDIS is sensitive to database coverage, query formulation and the selected component weights. Its value is comparative within a defined snapshot; it is not a clinical score. The explicit unavailable state and source links are intended to make these boundaries visible.")
    add_para(doc,"Several limitations define current development priorities. Bioactivity identifier reconciliation includes three unresolved historical symbols. API caches are process local. Some profile panels and association queries depend on live external services. Automated tests use mocks for MongoDB and scientific APIs and therefore complement, but do not replace, live integration checks. The production server must provide HTTPS access, help material, a licence link and sustained availability before NAR submission. A public code release and archived database snapshot should be deposited before submission to satisfy reproducibility and long-term availability requirements.")

    doc.add_heading("DATA AVAILABILITY", level=1)
    p=add_para(doc,"KinomeX is intended to be freely accessible without login at [PUBLIC HTTPS URL TO BE INSERTED]. Source code will be available at ")
    add_hyperlink(p,"https://github.com/dokhlab/kinomex","https://github.com/dokhlab/kinomex")
    p.add_run(" and should be archived with a versioned DOI before submission. The manuscript snapshot was audited on 11 August 2026. Original records remain available from UniProt, KinHub/KinMap, STRING, RCSB PDB, AlphaFold DB, ChEMBL, PubChem, GTEx, ClinVar, PubMed and ClinicalTrials.gov under each provider’s terms. KinomeX records preserve source identifiers and links. ChEMBL-derived redistribution remains subject to CC BY-SA 3.0; UniProt, STRING and AlphaFold DB data are CC BY 4.0; RCSB PDB data are provided under CC0. PubMed abstract copyright is publication specific. KinHub roster redistribution should be reviewed with its maintainers because an explicit current data licence was not identified. A complete attribution and reuse notice is included with the software.")

    doc.add_heading("SUPPLEMENTARY DATA", level=1)
    add_para(doc,"Supplementary Data are available at NAR Online. Proposed supplementary files include the catalogue-accounting audit, database-integrity audit, API examples, ETL configuration and the PDIS component specification.")
    doc.add_heading("ACKNOWLEDGEMENTS", level=1)
    add_para(doc,"The authors acknowledge the maintainers and contributors of the public resources listed in Table 1. [Additional acknowledgements to be inserted by the authors.]")
    doc.add_heading("AUTHOR CONTRIBUTIONS", level=1)
    add_para(doc,"[Author contribution assignments using the CRediT taxonomy are to be completed after the author list is finalized.]")
    doc.add_heading("FUNDING", level=1)
    add_para(doc,"[Funding agencies and grant numbers are to be inserted by the authors. Funding for the open-access charge: to be determined.]")
    doc.add_heading("CONFLICT OF INTEREST", level=1)
    add_para(doc,"The authors declare no conflict of interest. [Confirm before submission.]")

    doc.add_heading("REFERENCES", level=1)
    refs = [
        "1. Manning G, Whyte DB, Martinez R, Hunter T and Sudarsanam S. The protein kinase complement of the human genome. Science. 2002;298:1912–1934. doi:10.1126/science.1075762.",
        "2. Eid S, Turk S, Volkamer A, Rippmann F and Fulle S. KinMap: a web-based tool for interactive navigation through human kinome data. BMC Bioinformatics. 2017;18:16. doi:10.1186/s12859-016-1433-7.",
        "3. UniProt Consortium. UniProt: the Universal Protein Knowledgebase in 2025. Nucleic Acids Res. 2025;53:D609–D617. doi:10.1093/nar/gkae1010.",
        "4. Szklarczyk D, Kirsch R, Koutrouli M, Nastou K, Mehryary F, Hachilif R et al. The STRING database in 2023: protein–protein association networks and functional enrichment analyses for any sequenced genome of interest. Nucleic Acids Res. 2023;51:D638–D646. doi:10.1093/nar/gkac1000.",
        "5. Burley SK, Bhikadiya C, Bi C, Bittrich S, Chao H, Chen L et al. RCSB Protein Data Bank: tools for visualizing and understanding biological macromolecules in 3D. Nucleic Acids Res. 2025;53:D564–D574. doi:10.1093/nar/gkae1091.",
        "6. Varadi M, Anyango S, Deshpande M, Nair S, Natassia C, Yordanova G et al. AlphaFold Protein Structure Database: massively expanding the structural coverage of protein-sequence space with high-accuracy models. Nucleic Acids Res. 2022;50:D439–D444. doi:10.1093/nar/gkab1061.",
        "7. GTEx Consortium. The GTEx Consortium atlas of genetic regulatory effects across human tissues. Science. 2020;369:1318–1330. doi:10.1126/science.aaz1776.",
        "8. Landrum MJ, Lee JM, Benson M, Brown GR, Chao C, Chitipiralla S et al. ClinVar: improving access to variant interpretations and supporting evidence. Nucleic Acids Res. 2018;46:D1062–D1067. doi:10.1093/nar/gkx1153.",
        "9. Zdrazil B, Felix E, Hunter F, Manners EJ, Blackshaw J, Corbett S et al. The ChEMBL Database in 2023: a drug discovery platform spanning multiple bioactivity data types and time periods. Nucleic Acids Res. 2024;52:D1180–D1192. doi:10.1093/nar/gkad1004.",
        "10. Kim S, Chen J, Cheng T, Gindulyte A, He J, He S et al. PubChem 2025 update. Nucleic Acids Res. 2025;53:D1516–D1525. doi:10.1093/nar/gkae1059.",
        "11. Fiorini N, Lipman DJ and Lu Z. Towards PubMed 2.0. eLife. 2017;6:e28801. doi:10.7554/eLife.28801.",
        "12. Zarin DA, Tse T, Williams RJ and Carr S. Trial reporting in ClinicalTrials.gov—the final rule. N Engl J Med. 2016;375:1998–2004. doi:10.1056/NEJMsr1611785.",
        "13. Krokhotin A, Houlihan K and Dokholyan NV. iFoldRNA v2: folding RNA with constraints. Bioinformatics. 2015;31:2891–2893. doi:10.1093/bioinformatics/btv221.",
        "14. Ramachandran S, Kota P, Ding F and Dokholyan NV. Automated minimization of steric clashes in protein structures. Proteins. 2011;79:261–270. doi:10.1002/prot.22879.",
        "15. Dokholyan NV, Shakhnovich B and Shakhnovich EI. Expanding protein universe and its origin from the biological Big Bang. Proc Natl Acad Sci USA. 2002;99:14132–14136. doi:10.1073/pnas.202497999."
    ]
    for ref in refs: add_para(doc, ref, style="Reference")

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("AUTHOR COMPLETION CHECKLIST (REMOVE BEFORE SUBMISSION)", level=1)
    items = [
        "Finalize the author list, affiliations, correspondence email and CRediT contributions.",
        "Insert the public HTTPS server URL in the abstract and Data Availability statement.",
        "Confirm that the public server is free, requires no login, provides help/tutorial material and displays a licence link.",
        "Create a versioned public code release and archive the manuscript database snapshot with a DOI.",
        "Confirm all funding, open-access-charge funding and conflict-of-interest statements.",
        "Verify every bibliographic record against Crossref/PubMed and update access dates immediately before submission.",
        "Resolve or document the three historical bioactivity symbols and rerun both audits for the final snapshot.",
        "Prepare the NAR one-page Web Server proposal and submit by the applicable deadline before manuscript submission.",
        "Export the graphical abstract as a separate high-resolution file and run the NAR pre-submission checks."
    ]
    for item in items: add_para(doc,"☐ "+item)

    path = OUT / "KinomeX_NAR_manuscript_draft.docx"
    doc.save(path)
    print(path)


if __name__ == "__main__":
    build()
